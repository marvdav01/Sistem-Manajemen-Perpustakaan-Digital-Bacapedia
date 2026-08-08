from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# === PAGE MARGINS ===
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(3)
section.right_margin = Cm(2)

# === STYLES ===
style_normal = doc.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal.font.size = Pt(12)

# === HELPER FUNCTIONS ===
def set_font(run, bold=False, size=12, color=None, italic=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if level == 1:
        set_font(run, bold=True, size=16, color=(0, 70, 127))
    elif level == 2:
        set_font(run, bold=True, size=14, color=(21, 101, 192))
    elif level == 3:
        set_font(run, bold=True, size=12, color=(0, 0, 0))
    return p

def add_paragraph(text, bold=False, italic=False, indent=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hdr_cells[i].paragraphs[0].runs[0]
        set_font(run, bold=True, size=11)
        # Header background color (biru tua)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), '1565C0')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for ri, row_data in enumerate(rows):
        row_cells = table.rows[ri + 1].cells
        for ci, cell_text in enumerate(row_data):
            row_cells[ci].text = cell_text
            row_cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if ci > 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = row_cells[ci].paragraphs[0].runs[0]
            set_font(run, size=11)
            # Alternating row color
            if ri % 2 == 1:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:val'), 'clear')
                shading.set(qn('w:color'), 'auto')
                shading.set(qn('w:fill'), 'E3F2FD')
                row_cells[ci]._tc.get_or_add_tcPr().append(shading)

    # Set column widths if given
    if col_widths:
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[j])

    return table

def add_separator():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1565C0')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ============================
# COVER PAGE
# ============================
for _ in range(4):
    doc.add_paragraph()

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_title.add_run("DOKUMEN ERD DAN RANCANGAN SKEMA TABEL")
set_font(run, bold=True, size=20, color=(0, 70, 127))

doc.add_paragraph()

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_sub.add_run("Sistem Manajemen Perpustakaan Digital")
set_font(run, bold=True, size=16)

p_name = doc.add_paragraph()
p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_name.add_run('"Bacapedia"')
set_font(run, bold=True, size=18, color=(21, 101, 192))

for _ in range(2):
    doc.add_paragraph()

p_info = doc.add_paragraph()
p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_info.add_run("Diajukan untuk memenuhi Tugas Praktek / Demonstrasi\nSertifikasi Backend Developer - BNSP")
set_font(run, size=12, italic=True)

for _ in range(6):
    doc.add_paragraph()

p_year = doc.add_paragraph()
p_year.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_year.add_run("2026")
set_font(run, bold=True, size=14)

doc.add_page_break()

# ============================
# BAB 1: ERD
# ============================
add_heading("1. Entity Relationship Diagram (ERD)")
add_separator()
doc.add_paragraph()

add_paragraph(
    "Berikut adalah desain Entity Relationship Diagram (ERD) untuk sistem Bacapedia yang "
    "memvisualisasikan hubungan antar entitas utama beserta atribut-atribut yang dimilikinya.",
    indent=True
)
doc.add_paragraph()

# Relasi antar Entitas
add_heading("Relasi Antar Entitas", level=2)
relasi_data = [
    ("USERS", "PEMINJAMAN", "One-to-Many (1:N)", "Satu pengguna (Anggota) dapat melakukan banyak peminjaman."),
    ("KATEGORIS", "BUKUS", "One-to-Many (1:N)", "Satu kategori dapat memiliki banyak buku."),
    ("BUKUS", "PEMINJAMAN", "One-to-Many (1:N)", "Satu buku dapat dipinjam berkali-kali dalam transaksi berbeda."),
]

tbl_relasi = add_table(
    ["Entitas A", "Entitas B", "Tipe Relasi", "Keterangan"],
    relasi_data,
    col_widths=[3, 3, 4, 7]
)
doc.add_paragraph()

# Deskripsi ERD dalam bentuk teks
add_heading("Penjelasan Hubungan Entitas", level=2)

relasi_text = [
    ("USERS ⟶ PEMINJAMAN", "Seorang pengguna dengan role Anggota dapat melakukan satu hingga banyak transaksi peminjaman. Hubungan ini adalah One-to-Many."),
    ("KATEGORIS ⟶ BUKUS", "Setiap buku diklasifikasikan ke dalam satu kategori. Satu kategori dapat mengelompokkan banyak buku. Hubungan ini adalah One-to-Many."),
    ("BUKUS ⟶ PEMINJAMAN", "Sebuah buku dapat tercatat dalam banyak transaksi peminjaman oleh anggota yang berbeda. Hubungan ini adalah One-to-Many."),
]

for label, desc in relasi_text:
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(f"{label}: ")
    set_font(r1, bold=True, size=12)
    r2 = p.add_run(desc)
    set_font(r2, size=12)

doc.add_paragraph()

doc.add_page_break()

# ============================
# BAB 2: RANCANGAN SKEMA TABEL
# ============================
add_heading("2. Rancangan Skema Tabel (Data Dictionary)")
add_separator()
doc.add_paragraph()

add_paragraph(
    "Berikut adalah rincian detail dari masing-masing tabel pada basis data relasional (bacapedia_db) "
    "yang digunakan oleh sistem Bacapedia.",
    indent=True
)
doc.add_paragraph()

# === TABEL USERS ===
add_heading("A. Tabel users", level=2)
add_paragraph(
    "Tabel ini menyimpan seluruh data pengguna yang terdaftar dalam sistem, mencakup Admin, "
    "Petugas, dan Anggota. Hak akses dibedakan berdasarkan kolom role.",
    indent=True
)
doc.add_paragraph()

users_rows = [
    ("id", "BIGINT UNSIGNED", "Primary Key", "Auto Increment"),
    ("user_id", "VARCHAR(255)", "Unique", "UUID untuk identifikasi publik"),
    ("nama", "VARCHAR(255)", "-", "Nama lengkap pengguna"),
    ("email", "VARCHAR(255)", "Unique", "Email untuk keperluan login"),
    ("password", "VARCHAR(255)", "-", "Kata sandi yang sudah di-hash (Bcrypt)"),
    ("role", "ENUM", "Default: 'Anggota'", "Pilihan: 'Admin', 'Petugas', 'Anggota'"),
    ("created_at", "TIMESTAMP", "Nullable", "Waktu data dibuat"),
    ("updated_at", "TIMESTAMP", "Nullable", "Waktu data terakhir diubah"),
]
add_table(["Nama Kolom", "Tipe Data", "Constraint", "Keterangan"], users_rows, col_widths=[3.5, 3.5, 3.5, 6.5])
doc.add_paragraph()

# === TABEL KATEGORIS ===
add_heading("B. Tabel kategoris", level=2)
add_paragraph(
    "Tabel master kategori buku. Data di tabel ini hanya dapat dikelola (CRUD penuh) oleh pengguna "
    "dengan role Admin.",
    indent=True
)
doc.add_paragraph()

kategoris_rows = [
    ("id", "BIGINT UNSIGNED", "Primary Key", "Auto Increment"),
    ("nama_kategori", "VARCHAR(255)", "Unique", "Nama kategori (contoh: Fiksi, Sains)"),
    ("created_at", "TIMESTAMP", "Nullable", "Waktu data dibuat"),
    ("updated_at", "TIMESTAMP", "Nullable", "Waktu data terakhir diubah"),
]
add_table(["Nama Kolom", "Tipe Data", "Constraint", "Keterangan"], kategoris_rows, col_widths=[3.5, 3.5, 3.5, 6.5])
doc.add_paragraph()

# === TABEL BUKUS ===
add_heading("C. Tabel bukus", level=2)
add_paragraph(
    "Tabel ini merepresentasikan entitas koleksi buku dalam perpustakaan. Setiap buku "
    "berelasi dengan satu kategori pada tabel kategoris. Pengelolaan (tambah, ubah, hapus) "
    "hanya dapat dilakukan oleh Admin.",
    indent=True
)
doc.add_paragraph()

bukus_rows = [
    ("id", "BIGINT UNSIGNED", "Primary Key", "Auto Increment"),
    ("buku_id", "VARCHAR(255)", "Unique", "UUID identitas buku"),
    ("judul", "VARCHAR(255)", "-", "Judul lengkap buku"),
    ("penulis", "VARCHAR(255)", "-", "Nama penulis buku"),
    ("penerbit", "VARCHAR(255)", "-", "Nama penerbit buku"),
    ("kategori_id", "BIGINT UNSIGNED", "Foreign Key → kategoris.id", "Referensi ke tabel kategoris"),
    ("stok", "INT(11)", "Default: 0", "Jumlah ketersediaan buku"),
    ("tahun_terbit", "YEAR(4)", "-", "Tahun rilis penerbitan buku"),
    ("created_at", "TIMESTAMP", "Nullable", "Waktu data dibuat"),
    ("updated_at", "TIMESTAMP", "Nullable", "Waktu data terakhir diubah"),
]
add_table(["Nama Kolom", "Tipe Data", "Constraint", "Keterangan"], bukus_rows, col_widths=[3.5, 3.5, 4, 6])
doc.add_paragraph()

doc.add_page_break()

# === TABEL PEMINJAMEN ===
add_heading("D. Tabel peminjamen (Peminjaman)", level=2)
add_paragraph(
    "Tabel transaksional yang mencatat seluruh riwayat pergerakan peminjaman dan "
    "pengembalian buku. Tabel ini berelasi dengan tabel users (Anggota yang meminjam) "
    "dan tabel bukus (buku yang dipinjam).",
    indent=True
)
doc.add_paragraph()

peminjamen_rows = [
    ("id", "BIGINT UNSIGNED", "Primary Key", "Auto Increment"),
    ("user_id", "BIGINT UNSIGNED", "Foreign Key → users.id", "ID Anggota yang meminjam"),
    ("buku_id", "BIGINT UNSIGNED", "Foreign Key → bukus.id", "ID Buku yang dipinjam"),
    ("tanggal_pinjam", "DATE", "-", "Tanggal disetujuinya peminjaman"),
    ("tanggal_jatuh_tempo", "DATE", "-", "Batas pengembalian (7 hari dari pinjam)"),
    ("tanggal_kembali", "DATE", "Nullable", "Diisi otomatis saat buku dikembalikan"),
    ("status", "ENUM", "Default: 'dipinjam'", "Pilihan: 'dipinjam', 'dikembalikan'"),
    ("denda", "INT(11)", "Default: 0", "Nominal denda (Rp 2.000/hari keterlambatan)"),
    ("created_at", "TIMESTAMP", "Nullable", "Waktu data dibuat"),
    ("updated_at", "TIMESTAMP", "Nullable", "Waktu data terakhir diubah"),
]
add_table(["Nama Kolom", "Tipe Data", "Constraint", "Keterangan"], peminjamen_rows, col_widths=[3.5, 3.5, 4, 6])
doc.add_paragraph()

# ============================
# BAB 3: CATATAN RELASIONAL
# ============================
add_heading("3. Catatan Logika Relasional")
add_separator()
doc.add_paragraph()

notes = [
    "Ketika sebuah Kategori dihapus, seluruh Buku yang menggunakan kategori tersebut akan ikut terhapus secara otomatis (ON DELETE CASCADE).",
    "Ketika sebuah data Buku dihapus dari sistem, seluruh riwayat Peminjaman yang terkait buku tersebut juga akan ikut terhapus demi menjaga integritas data (ON DELETE CASCADE).",
    "Ketika sebuah data User (Anggota) dihapus, seluruh riwayat Peminjaman yang terkait pengguna tersebut juga akan ikut terhapus secara otomatis (ON DELETE CASCADE).",
    "Kolom stok pada tabel bukus akan berkurang (-1) secara otomatis ketika terjadi peminjaman, dan akan bertambah (+1) saat buku berhasil dikembalikan.",
    "Kolom password pada tabel users selalu disimpan dalam format Hash (Bcrypt) dan tidak pernah disimpan sebagai teks biasa (plain text).",
]

for note in notes:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(note)
    set_font(run, size=12)

doc.add_paragraph()

# Footer info
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_footer.add_run("— Dokumen ini dibuat sebagai bagian dari Deliverables Asesmen Sertifikasi Backend Developer BNSP —")
set_font(run, size=10, italic=True, color=(100, 100, 100))

# ============================
# SAVE
# ============================
output_path = r"C:\Users\MyBook Hype AMD\Documents\Sistem-Manajemen-Perpustakaan-Digital-Bacapedia-\Dokumen_ERD_Bacapedia.docx"
doc.save(output_path)
print("SUKSES! File Word tersimpan di: " + output_path)
